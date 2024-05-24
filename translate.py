import os
import re
import time
import PyPDF2
from ebooklib import epub
import ebooklib
from docx import Document
from google.cloud import translate_v2 as translate

def count_words(text):
    words = re.findall(r'\w+', text)
    return len(words)

def translate_text(text, target_language):
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "/home/hp-laserjet-wrt8070/Desktop/tidal-memento-421014-9813bcb0b561.json"
    translator = translate.Client()
    
    lines = text.split('\n')
    translated_lines = []

    total_lines = len(lines)
    translated_lines_count = 0
    estimated_remaining_time = 0
    start_time = time.time()

    for line in lines:
        result = translator.translate(line, target_language=target_language)
        translated_line = result['translatedText']
        translated_lines.append(translated_line)

        translated_lines_count += 1
        completion_percentage = (translated_lines_count / total_lines) * 100

        current_time = time.time()
        elapsed_time = current_time - start_time
        if translated_lines_count > 0:
            estimated_remaining_time = (elapsed_time / translated_lines_count) * (total_lines - translated_lines_count)

        print(f"Tamamlanan yüzde: {completion_percentage:.2f}%, Tahmini kalan süre: {estimated_remaining_time:.2f} saniye")

    translated_text = '\n'.join(translated_lines)
    total_time = time.time() - start_time
    print(f"Toplam geçen süre: {total_time:.2f} saniye")

    translated_word_count = count_words(translated_text)
    return translated_text, translated_word_count

def translate_pdf(input_file, output_file, target_language):
    pdf_text = ''
    with open(input_file, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        writer = PyPDF2.PdfWriter()
        
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            page_text = page.extract_text()
            translated_text, _ = translate_text(page_text, target_language)
            page.text = translated_text
            writer.add_page(page)

        with open(output_file, 'wb') as output:
            writer.write(output)

    print(f"PDF dosyası başarıyla çevrildi: {output_file}")

def translate_docx(input_file, output_file, target_language):
    doc = Document(input_file)
    translated_doc = Document()
    
    for paragraph in doc.paragraphs:
        translated_text, _ = translate_text(paragraph.text, target_language)
        translated_doc.add_paragraph(translated_text)
    
    translated_doc.save(output_file)
    print(f"DOCX dosyası başarıyla çevrildi: {output_file}")

def translate_epub(input_file, output_file, target_language):
    book = epub.read_epub(input_file)
    translated_book = epub.EpubBook()
    
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            content = item.get_content()
            translated_text, _ = translate_text(content.decode('utf-8'), target_language)
            item.content = translated_text.encode('utf-8')
            translated_book.add_item(item)

    translated_book.set_title(book.get_title())
    translated_book.set_language(target_language)
    translated_book.set_identifier(book.get_identifier())

    epub.write_epub(output_file, translated_book)
    print(f"EPUB dosyası başarıyla çevrildi: {output_file}")

def translate_txt(input_file, output_file, target_language):
    with open(input_file, 'r', encoding='utf-8') as file:
        text = file.read()
    translated_text, _ = translate_text(text, target_language)
    
    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(translated_text)
    
    print(f"TXT dosyası başarıyla çevrildi: {output_file}")



input_file_pdf = "/home/hp-laserjet-wrt8070/Desktop/test_iktidar_pdf.pdf"
output_file_pdf = "/home/hp-laserjet-wrt8070/Desktop/translated_test_iktidar_pdf.pdf"

target_language = "tr"


translate_pdf(input_file_pdf, output_file_pdf, target_language)
